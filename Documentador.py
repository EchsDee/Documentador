import os
import tempfile
from io import BytesIO
from flask import Flask, render_template, request, send_file, after_this_request
from docx import Document
from docx.shared import Inches
import requests
import json
from apscheduler.schedulers.background import BackgroundScheduler
from werkzeug.utils import secure_filename
from docx.opc.exceptions import PackageNotFoundError
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging
from docx.oxml.ns import qn
from waitress import serve

app = Flask(__name__, static_folder='static')

template_path = os.path.join(os.getcwd(),'doc', 'TemplateDocument.docx')
api_url = os.environ.get('API_URL') or 'http://3.140.207.100/api/getclientes.php'

temp_dir = tempfile.mkdtemp()
half_template_path = os.path.join(os.getcwd(),'doc', 'templateDocument-half-segundo.docx')  
scheduler = BackgroundScheduler()
scheduler.add_job(lambda: clean_temp_folder(temp_dir), 'interval', minutes=5)
scheduler.start()

def clean_temp_folder(temp_folder):
    for file_name in os.listdir(temp_folder):
        file_path = os.path.join(temp_folder, file_name)
        try:
            if os.path.isfile(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"Error deleting file: {e}")

# Function to fetch data from the API
def get_data_from_api(api_url):
    try:
        response = requests.get(api_url)
        response.raise_for_status()  # Raise HTTPError for bad responses
        content = response.content.decode('utf-8-sig')  # Decode using utf-8-sig to handle BOM
        return json.loads(content)
    except requests.RequestException as e:
        print(f"Error fetching data from API: {e}")
        return []
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")
        return []

@app.route('/doc')
def doc():
    clients_data = get_data_from_api(api_url)
    return render_template('doc.html', clients=clients_data)

@app.route('/teste')
def teste():
    clients_data = get_data_from_api(api_url)
    return render_template('teste.html', clients=clients_data)

@app.route('/')
def front():
    return render_template('front.html')

@app.route('/process_template', methods=['POST'])
def process_template():
    # Get form data
    data1 = request.form.get('data1', '')
    data2 = request.form.get('data2', '')
    data3 = request.form.get('data3', '')
    data4 = request.form.get('data4', '')
    data5 = request.form.get('data5', '')
    usuario = request.form.get('hiddenUsuario', 'nao informado')
    support_level = request.form.get('supportLevel', '')

    # Prepare image files and descriptions
    image_files = request.files.getlist('data6[]')
    image_descriptions = request.form.getlist('data7[]')

    # Set modified document path
    if 'additionalFile' in request.files:
        # Handle uploaded document
        uploaded_file = request.files['additionalFile']
        uploaded_filename = secure_filename(uploaded_file.filename)
        modified_filename = f"{os.path.splitext(uploaded_filename)[0]} {usuario} - testado.docx"
        modified_path = os.path.join(temp_dir, modified_filename)
        additional_file_path = os.path.join(temp_dir, uploaded_filename)
        uploaded_file.save(additional_file_path)

        try:
            # Merge uploaded document with half-template
            process_uploaded_doc(additional_file_path, half_template_path, modified_path)
        except PackageNotFoundError:
            return "Uploaded file is not a valid DOCX file.", 400
        # Define 'doc' after processing
        doc = Document(modified_path)
    else:
        # No additional file uploaded, use default values
        modified_filename = f'DOCUMENTAÇÃO - {secure_filename(data2 or "document")}.docx'
        modified_path = os.path.join(temp_dir, modified_filename)
        doc = Document(template_path)
        doc.save(modified_path)

    # Now 'doc' is defined in both cases
    # Replace placeholders in the document
    replace_placeholder(doc, '@chamado', data1)
    replace_placeholder(doc, '@cliente', data2)
    replace_placeholder(doc, '@modulo', data3)
    replace_placeholder(doc, '@data', data4)
    replace_placeholder(doc, '@descricao', data5)
    replace_placeholder(doc, '@usuario', usuario)
    doc.save(modified_path)

    # Insert images into the document
    insert_all_images_with_description(modified_path, image_files, image_descriptions)

    return send_file(modified_path, as_attachment=True, download_name=modified_filename)

def replace_placeholder(doc, placeholder, replacement):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)

def save_image(file):
    image_folder = temp_dir  # Use the global temp_dir
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)

    image_path = os.path.join(image_folder, secure_filename(file.filename))
    file.save(image_path)
    return image_path

def insert_all_images_with_description(doc_path, image_files, image_descriptions, insertion_placeholder='@prints'):
    doc = Document(doc_path)
    
    # Flag to check if placeholder is found
    placeholder_found = False
    
    # Search in body paragraphs
    for idx, paragraph in enumerate(doc.paragraphs):
        if insertion_placeholder in paragraph.text:
            placeholder_found = True
            # Remove the placeholder paragraph
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None

            # Insert images and descriptions at the placeholder position
            for image_file, description in zip(image_files, image_descriptions):
                # Add description as a new paragraph
                description_paragraph = doc.add_paragraph(description)
                description_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.paragraphs.insert(idx, description_paragraph)
                idx += 1
                # Add image as a new paragraph with increased size
                image_path = save_image(image_file)
                picture_paragraph = doc.add_paragraph()
                picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                picture_run = picture_paragraph.add_run()
                picture_run.add_picture(image_path, width=Inches(2.0))  # Increased width
                doc.paragraphs.insert(idx, picture_paragraph)
                idx += 1
                # Add an empty paragraph for spacing
                spacing_paragraph = doc.add_paragraph()
                spacing_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                doc.paragraphs.insert(idx, spacing_paragraph)
                idx += 1
            break  # Placeholder found and processed
    
    # If placeholder was not found in body, search in table cells
    if not placeholder_found:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for idx, paragraph in enumerate(cell.paragraphs):
                        if insertion_placeholder in paragraph.text:
                            # Remove the placeholder paragraph
                            p = paragraph._element
                            p.getparent().remove(p)
                            p._p = p._element = None

                            # Insert images and descriptions at the placeholder position
                            for image_file, description in zip(image_files, image_descriptions):
                                # Add description as a new paragraph
                                description_paragraph = cell.add_paragraph(description)
                                description_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                cell.paragraphs.insert(idx, description_paragraph)
                                idx += 1
                                # Add image as a new paragraph with increased size
                                image_path = save_image(image_file)
                                picture_paragraph = cell.add_paragraph()
                                picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                picture_run = picture_paragraph.add_run()
                                picture_run.add_picture(image_path, width=Inches(5.0))  # Increased width
                                cell.paragraphs.insert(idx, picture_paragraph)
                                idx += 1
                                # Add an empty paragraph for spacing
                                spacing_paragraph = cell.add_paragraph()
                                spacing_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                cell.paragraphs.insert(idx, spacing_paragraph)
                                idx += 1
                            break  # Placeholder found and processed
                    else:
                        continue
                    break
                else:
                    continue
                break
            else:
                continue
            break

    doc.save(doc_path)


def delete_between_markers(doc, start_text, end_text):
    body = doc.element.body
    start_found = False
    elements_to_remove = []

    for child in list(body):
        if not start_found:
            if child.tag == qn('w:p'):
                paragraphs = child.findall('.//w:t', namespaces=child.nsmap)
                para_text = ''.join([t.text for t in paragraphs if t.text])
                if start_text in para_text:
                    logging.debug(f"Start marker '{start_text}' found.")
                    start_found = True
            continue
        else:
            if child.tag == qn('w:p'):
                paragraphs = child.findall('.//w:t', namespaces=child.nsmap)
                para_text = ''.join([t.text for t in paragraphs if t.text])
                if end_text in para_text:
                    logging.debug(f"End marker '{end_text}' found. Removing this element.")
                    elements_to_remove.append(child)
                    break
            elements_to_remove.append(child)

    for element in elements_to_remove:
        logging.debug(f"Removing element: {element.tag}")
        body.remove(element)

def process_uploaded_doc(uploaded_doc_path, half_template_path, output_path):
    # Open the uploaded document
    uploaded_doc = Document(uploaded_doc_path)
    # Open the half-template document
    half_template_doc = Document(half_template_path)
    
    
    # Replace placeholders in both documents
    placeholders = {
        '@chamado': request.form.get('data1', ''),
        '@cliente': request.form.get('data2', ''),
        '@modulo': request.form.get('data3', ''),
        '@data': request.form.get('data4', ''),
        '@descricao': request.form.get('data5', ''),
        '@usuario': request.form.get('hiddenUsuario', 'nao informado')
    }

    for placeholder, value in placeholders.items():
        replace_placeholder(uploaded_doc, placeholder, value)
        replace_placeholder(half_template_doc, placeholder, value)
      
    # Merge the modified half-template into the uploaded document
    for element in half_template_doc.element.body:
        uploaded_doc.element.body.append(element)
    
    # Delete content between markers
    delete_between_markers(uploaded_doc, "PREENCHIMENTO DO TESTE E QUALIDADE", "@stop")
    
    # Save the merged document
    uploaded_doc.save(output_path)

#if __name__ == '__main__':
   #app.run(host='localhost', port=8000, debug=True)

#Modificado operação para Waitress pois o Flask não é recomendado para produção 
if __name__ == '__main__':
    serve(app, host='localhost', port=8000)
